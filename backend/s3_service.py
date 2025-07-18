import boto3
import os
import logging
from typing import Optional, List, Dict, Any, Tuple
from botocore.exceptions import ClientError
import PyPDF2
import docx
import io
from datetime import datetime
import mimetypes

logger = logging.getLogger(__name__)

class S3FileService:
    """AWS S3 service for file storage and processing"""
    
    def __init__(self):
        self.s3_client = boto3.client(
            's3',
            region_name=os.getenv('AWS_REGION', 'us-east-1'),
            aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
            aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY')
        )
        
        self.bucket_name = os.getenv('S3_BUCKET_NAME')
        if not self.bucket_name:
            raise ValueError("S3_BUCKET_NAME environment variable is required")
    
    async def upload_file(
        self, 
        file_content: bytes, 
        file_name: str, 
        user_id: str,
        content_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """Upload file to S3 and return file info"""
        try:
            # Generate unique file key
            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
            file_extension = os.path.splitext(file_name)[1]
            unique_filename = f"{timestamp}_{file_name}"
            file_key = f"users/{user_id}/files/{unique_filename}"
            
            # Detect content type if not provided
            if not content_type:
                content_type, _ = mimetypes.guess_type(file_name)
                if not content_type:
                    content_type = 'application/octet-stream'
            
            # Upload to S3
            self.s3_client.put_object(
                Bucket=self.bucket_name,
                Key=file_key,
                Body=file_content,
                ContentType=content_type,
                Metadata={
                    'user_id': user_id,
                    'original_name': file_name,
                    'upload_timestamp': timestamp
                }
            )
            
            # Generate file URL
            file_url = f"https://{self.bucket_name}.s3.{os.getenv('AWS_REGION', 'us-east-1')}.amazonaws.com/{file_key}"
            
            return {
                'file_key': file_key,
                'file_url': file_url,
                'file_size': len(file_content),
                'content_type': content_type,
                'original_name': file_name
            }
            
        except ClientError as e:
            logger.error(f"S3 upload error: {e}")
            raise Exception(f"Failed to upload file: {e}")
        except Exception as e:
            logger.error(f"File upload error: {e}")
            raise Exception(f"Upload failed: {e}")
    
    async def get_file(self, file_key: str) -> bytes:
        """Download file from S3"""
        try:
            response = self.s3_client.get_object(
                Bucket=self.bucket_name,
                Key=file_key
            )
            return response['Body'].read()
        except ClientError as e:
            logger.error(f"S3 download error: {e}")
            raise Exception(f"Failed to download file: {e}")
    
    async def delete_file(self, file_key: str) -> bool:
        """Delete file from S3"""
        try:
            self.s3_client.delete_object(
                Bucket=self.bucket_name,
                Key=file_key
            )
            return True
        except ClientError as e:
            logger.error(f"S3 delete error: {e}")
            return False
    
    async def extract_text_from_file(
        self, 
        file_content: bytes, 
        file_name: str,
        content_type: str
    ) -> Tuple[str, bool]:
        """Extract text content from uploaded files"""
        try:
            extracted_text = ""
            success = False
            
            file_extension = os.path.splitext(file_name)[1].lower()
            
            if file_extension == '.pdf' or content_type == 'application/pdf':
                extracted_text, success = self._extract_pdf_text(file_content)
                
            elif file_extension in ['.docx', '.doc'] or content_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]:
                extracted_text, success = self._extract_docx_text(file_content)
                
            elif file_extension == '.txt' or content_type == 'text/plain':
                try:
                    raw_text = file_content.decode('utf-8', errors='ignore')
                    extracted_text = self._clean_extracted_text(raw_text)
                    success = True
                except Exception as e:
                    logger.error(f"Text file decoding error: {e}")
                    extracted_text = f"Error reading text file: {str(e)}"
                    success = False
                
            else:
                logger.warning(f"Unsupported file type: {file_extension}, {content_type}")
                extracted_text = f"File type {file_extension} is not supported for text extraction."
                success = False
            
            # Final safety check - ensure the text is safe for database storage
            if success and extracted_text:
                try:
                    # Test if the text can be safely encoded/decoded
                    test_encoded = extracted_text.encode('utf-8')
                    test_decoded = test_encoded.decode('utf-8')
                    
                    # Limit the text length to prevent database issues
                    if len(extracted_text) > 100000:  # 100KB limit
                        extracted_text = extracted_text[:100000] + "\n\n[Text truncated due to length...]"
                        logger.info(f"Text truncated for {file_name} (original length: {len(extracted_text)})")
                    
                except UnicodeError as e:
                    logger.error(f"Unicode error in extracted text from {file_name}: {e}")
                    extracted_text = "Error: File contains characters that cannot be processed"
                    success = False
            
            return extracted_text, success
            
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return f"Error extracting text from {file_name}: {str(e)}", False
    
    def _extract_pdf_text(self, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        # Clean the text to handle encoding issues
                        cleaned_text = self._clean_extracted_text(page_text)
                        if cleaned_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{cleaned_text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
            
            if text_content:
                final_text = "\n\n".join(text_content)
                return final_text, True
            else:
                return "No text content found in PDF", False
                
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return f"Error reading PDF: {str(e)}", False
    
    def _extract_docx_text(self, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    cleaned_text = self._clean_extracted_text(paragraph.text)
                    if cleaned_text.strip():
                        text_content.append(cleaned_text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            cleaned_text = self._clean_extracted_text(cell.text)
                            if cleaned_text.strip():
                                row_text.append(cleaned_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if text_content:
                final_text = "\n\n".join(text_content)
                return final_text, True
            else:
                return "No text content found in document", False
                
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return f"Error reading DOCX: {str(e)}", False
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean extracted text to handle encoding issues and problematic characters"""
        try:
            import re
            
            # Remove or replace problematic Unicode characters
            # Remove surrogate pairs and other problematic characters
            text = re.sub(r'[\ud800-\udfff]', '', text)  # Remove surrogate pairs
            text = re.sub(r'[\ufeff\ufffe]', '', text)   # Remove BOM characters
            text = re.sub(r'[\u0000-\u0008\u000b\u000c\u000e-\u001f]', '', text)  # Remove control characters
            
            # Replace common problematic characters with ASCII equivalents
            replacements = {
                '\u2018': "'",  # Left single quotation mark
                '\u2019': "'",  # Right single quotation mark
                '\u201c': '"',  # Left double quotation mark
                '\u201d': '"',  # Right double quotation mark
                '\u2013': '-',  # En dash
                '\u2014': '-',  # Em dash
                '\u2026': '...',  # Horizontal ellipsis
                '\u00a0': ' ',  # Non-breaking space
            }
            
            for old_char, new_char in replacements.items():
                text = text.replace(old_char, new_char)
            
            # Normalize whitespace
            text = re.sub(r'\s+', ' ', text)  # Replace multiple whitespace with single space
            text = text.strip()
            
            # Ensure the final text is valid UTF-8
            # This will replace any remaining problematic characters
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            
            return text
            
        except Exception as e:
            logger.error(f"Text cleaning error: {e}")
            # If cleaning fails, try basic cleanup
            try:
                return text.encode('utf-8', errors='ignore').decode('utf-8').strip()
            except:
                return "Error processing text content"
    
    async def generate_presigned_url(
        self, 
        file_key: str, 
        expiration: int = 3600
    ) -> str:
        """Generate a presigned URL for file access"""
        try:
            response = self.s3_client.generate_presigned_url(
                'get_object',
                Params={'Bucket': self.bucket_name, 'Key': file_key},
                ExpiresIn=expiration
            )
            return response
        except ClientError as e:
            logger.error(f"Presigned URL error: {e}")
            raise Exception(f"Failed to generate access URL: {e}")

# Create a singleton instance
s3_service = S3FileService()
